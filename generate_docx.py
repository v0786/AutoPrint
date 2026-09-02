import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def create_system_doc():
    doc = docx.Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("AutoPrint Enterprise System\nComprehensive Technical Documentation")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Architecture • Installation • Operations • Guidelines\nVersion 1.0.0 — Production Release")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph() # spacing

    # 1. System Overview
    h1 = doc.add_heading("1. System Overview & Objectives", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    doc.add_paragraph(
        "AutoPrint is an automated, fail-safe print collection and merchant management platform designed "
        "for print shops, retail counters, and self-service document kiosks. The system streamlines document "
        "submission, watermarking, verification code generation, payment processing, and document handover."
    )

    # Key Features Bullet List
    doc.add_paragraph("Key System Capabilities:", style='List Bullet')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("8-Digit Secure Verification Code: ").bold = True
    p.add_run("Cryptographically random 8-digit key (e.g. 4829 1057) generated per job with SHA-256 integrity checksum.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Automated Watermarking Overlay: ").bold = True
    p.add_run("Stamps verification key, checksum, and timestamp on the final page footer without disrupting page margins.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Fail-Safe 3-Strike Lockout Engine: ").bold = True
    p.add_run("Automatically locks digital payment options after 3 consecutive failures, enforcing mandatory counter cash collection.")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Background Windows Service: ").bold = True
    p.add_run("Runs invisibly on configured ports with automated browser portal launching.")

    # 2. Technical Architecture
    h2 = doc.add_heading("2. System Architecture & Folder Layout", level=1)
    h2.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    doc.add_paragraph(
        "The system enforces strict separation of concerns across backend REST services, "
        "merchant desktop terminals, and customer web applications."
    )

    # Architecture Table
    table = doc.add_table(rows=5, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    headers = ["Component Module", "Technology Stack", "Primary Responsibility"]
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(hdr_cells[i], "003366")

    row_data = [
        ("Backend Service", "Node.js, Express, TypeScript", "API routing, verification logic, 3-strike lockout, audit logs"),
        ("Merchant Desktop UI", "React, Vite, Electron Bridge", "Staff verification lookup, queue monitoring, cash reconciliation"),
        ("Customer Web UI", "React, Vite, TailwindCSS", "Document upload, print spec selection, UPI payment, code display"),
        ("Database Layer", "In-Memory / Persistent JSON Store", "Storage of active jobs, verification records, and audit events")
    ]

    for row_idx, data in enumerate(row_data):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "F2F4F7")

    doc.add_paragraph()

    # 3. Hardware & Software Requirements
    h3 = doc.add_heading("3. Hardware & Software Requirements", level=1)
    h3.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    req_table = doc.add_table(rows=5, cols=2)
    req_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    req_table.style = 'Table Grid'
    
    r_headers = ["Resource Category", "Minimum Requirement"]
    for i, title in enumerate(r_headers):
        req_table.rows[0].cells[i].text = title
        req_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        req_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(req_table.rows[0].cells[i], "003366")

    req_data = [
        ("Operating System", "Windows 10 / Windows 11 (64-bit)"),
        ("Runtime Engine", "Node.js v18.0+ LTS or v20.0+ LTS"),
        ("User Privileges", "Administrator permissions (required for UAC folder setup)"),
        ("Network Ports", "Port 5000 (Backend), Port 3000 (Customer UI), Port 3001 (Merchant UI)")
    ]

    for r_idx, r_item in enumerate(req_data):
        r_cells = req_table.rows[r_idx + 1].cells
        r_cells[0].text = r_item[0]
        r_cells[1].text = r_item[1]
        if r_idx % 2 == 1:
            set_cell_background(r_cells[0], "F2F4F7")
            set_cell_background(r_cells[1], "F2F4F7")

    doc.add_paragraph()

    # 4. Installation Guide
    h4 = doc.add_heading("4. Step-by-Step Installation Guide", level=1)
    h4.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    doc.add_paragraph("Option 1: One-Click Windows Automated Setup", style='List Number')
    doc.add_paragraph("1. Extract the downloaded release archive onto the target machine.", style='List Bullet 2')
    doc.add_paragraph("2. Right-click installer.bat and choose 'Run as Administrator'.", style='List Bullet 2')
    doc.add_paragraph("3. The installer provisions C:\\AutoPrint and launches background services automatically.", style='List Bullet 2')

    doc.add_paragraph("Option 2: Executable Wizard Installer (.exe)", style='List Number')
    doc.add_paragraph("1. Compile setup.iss using Inno Setup (iscc setup.iss).", style='List Bullet 2')
    doc.add_paragraph("2. Run the generated AutoPrint_Setup_v1.0.0.exe installer.", style='List Bullet 2')

    # 5. What to Do & What NOT to Do
    h5 = doc.add_heading("5. Operational Guidelines: What to Do & What NOT to Do", level=1)
    h5.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    do_table = doc.add_table(rows=6, cols=2)
    do_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    do_table.style = 'Table Grid'
    
    do_table.rows[0].cells[0].text = "WHAT TO DO (BEST PRACTICES)"
    do_table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
    do_table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_background(do_table.rows[0].cells[0], "2E7D32")

    do_table.rows[0].cells[1].text = "WHAT NOT TO DO (CRITICAL RESTRICTIONS)"
    do_table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
    do_table.rows[0].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_background(do_table.rows[0].cells[1], "C62828")

    guidelines = [
        ("Always verify the physical printed document against the 8-digit verification code.", "DO NOT hand over printed documents before confirming payment or completing cash collection."),
        ("Ensure installer.bat is executed with Administrator privileges for directory binding.", "DO NOT modify or delete runtime database files inside C:\\AutoPrint\\Data directly."),
        ("Log counter staff interactions using assigned cashier credentials for audit compliance.", "DO NOT bypass the 3-strike cash lockout policy for customers with repeated digital failures."),
        ("Review daily audit logs in C:\\AutoPrint\\Logs\\ for transaction reconciliation.", "DO NOT change default network ports without updating appsettings.json across all services."),
        ("Keep backup snapshots of C:\\AutoPrint\\Config\\ for rapid disaster recovery.", "DO NOT run multiple instances of the backend service on conflicting ports.")
    ]

    for g_idx, (do_txt, dont_txt) in enumerate(guidelines):
        g_cells = do_table.rows[g_idx + 1].cells
        g_cells[0].text = f"✓ {do_txt}"
        g_cells[1].text = f"✗ {dont_txt}"
        if g_idx % 2 == 1:
            set_cell_background(g_cells[0], "F1F8E9")
            set_cell_background(g_cells[1], "FFEBEE")

    os.makedirs("E:/Project/AutoPrint/DOC", exist_ok=True)
    doc.save("E:/Project/AutoPrint/DOC/AutoPrint_System_Documentation.docx")
    print("Created AutoPrint_System_Documentation.docx successfully.")

def create_sdlc_doc():
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("AutoPrint Enterprise System\nSoftware Development Life Cycle (SDLC) Specification")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Requirements • Design • Implementation • QA • Maintenance\nStandard Compliance Document")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # 1. SDLC Methodology
    h1 = doc.add_heading("1. SDLC Methodology & Fail-Safe Engineering", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    doc.add_paragraph(
        "The AutoPrint development workflow adheres to an iterative Agile-Waterframe model structured "
        "around fail-safe defensive engineering principles. Every transactional phase incorporates deterministic "
        "verification state machines, cryptographically secure code generation, and immutable audit logging."
    )

    # 2. Phase 1: Requirements Analysis
    h2 = doc.add_heading("2. Phase 1: Requirements Analysis & Specifications", level=1)
    h2.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    doc.add_paragraph("Functional Requirements:", style='List Bullet')
    p = doc.add_paragraph(style='List Bullet 2')
    p.add_run("FR-1 Verification Code Generator: ").bold = True
    p.add_run("Generates 8-digit random codes with SHA-256 HMAC checksums.")

    p = doc.add_paragraph(style='List Bullet 2')
    p.add_run("FR-2 Final Page Watermarking: ").bold = True
    p.add_run("Appends non-intrusive stamp containing code and checksum on document footers.")

    p = doc.add_paragraph(style='List Bullet 2')
    p.add_run("FR-3 Fail-Safe 3-Strike Lockout: ").bold = True
    p.add_run("Locks payment method to counter cash collection after 3 failed digital attempts.")

    # 3. Phase 2: System Architecture Design
    h3 = doc.add_heading("3. Phase 2: Architecture & Data Modeling", level=1)
    h3.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    doc.add_paragraph(
        "The backend verification engine maintains state machine transitions across PENDING, "
        "UPI_SUCCESS, UPI_FAILED, CASH_LOCKED, CASH_REQUIRED, CASH_COLLECTED, and HANDOVER_COLLECTED."
    )

    # 4. Phase 3: Implementation & Defensive Coding
    h4 = doc.add_heading("4. Phase 3: Implementation & Security Standards", level=1)
    h4.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    doc.add_paragraph(
        "Implementation relies on TypeScript strict type definitions, Web Cryptography API random sampling, "
        "and clean separation of business logic across controllers, services, and audit loggers."
    )

    # 5. Phase 4: Quality Assurance & Test Strategy
    h5 = doc.add_heading("5. Phase 4: Quality Assurance & Testing Results", level=1)
    h5.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    doc.add_paragraph(
        "Empirical automated test suites (test_runner.ts) execute end-to-end simulation of job submissions, "
        "code generation, 3-strike failure lockouts, counter cash collection, and audit event verification."
    )

    # 6. Phase 5: Deployment & Maintenance
    h6 = doc.add_heading("6. Phase 5: Deployment & Operational Maintenance", level=1)
    h6.runs[0].font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    doc.add_paragraph(
        "Production deployment utilizes automated Windows installer scripts (installer.bat, install.ps1) "
        "and Inno Setup executable wrappers (setup.iss) for background service setup."
    )

    doc.save("E:/Project/AutoPrint/DOC/AutoPrint_SDLC_Documentation.docx")
    print("Created AutoPrint_SDLC_Documentation.docx successfully.")

if __name__ == "__main__":
    create_system_doc()
    create_sdlc_doc()
